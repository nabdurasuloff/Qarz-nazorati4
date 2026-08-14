# -*- coding: utf-8 -*-
"""
Word shabloniga ma'lumotlarni joylab, tayyor xat (.docx) yaratish.
"""
import os
import re
import sys
import datetime
from docx import Document


def _base_dir():
    # PyInstaller --onefile bilan yig'ilganda fayllar vaqtinchalik papkaga
    # (sys._MEIPASS) ochiladi; oddiy ishga tushirishda esa shu faylning papkasi.
    if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
        return sys._MEIPASS
    return os.path.dirname(os.path.abspath(__file__))


TEMPLATE_PATH = os.path.join(_base_dir(), 'templates', 'xat_shablon.docx')


def _fmt_summa(val):
    try:
        val = float(val)
    except (TypeError, ValueError):
        return str(val)
    return f"{val:,.0f}".replace(',', ' ')


def _replace_in_paragraph(paragraph, mapping):
    full_text = ''.join(run.text for run in paragraph.runs)
    if '{{' not in full_text:
        return
    new_text = full_text
    for key, val in mapping.items():
        new_text = new_text.replace('{{' + key + '}}', str(val))
    if new_text == full_text:
        return
    # Barcha runlarni tozalab, birinchi runga yangi matnni yozamiz
    # (formatlashni birinchi run’dan olamiz)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(new_text)


def _replace_everywhere(doc, mapping):
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)


def generate_letter(output_path, xat_turi, mijoz_ism, mijoz_manzil, portfel_row, settings,
                     anketa_raqami=None, rahbar_ism=None):
    """
    xat_turi: 'Ogohlantirish' yoki 'Talabnoma'
    portfel_row: dict — bitta portfel qatoridagi ma'lumot (database.py formatida)
    settings: dict — get_all_settings() natijasi
    """
    doc = Document(TEMPLATE_PATH)

    sarlavha = 'ОГОҲЛАНТИРИШ ХАТИ' if xat_turi == 'Ogohlantirish' else 'ТАЛАБНОМА'

    holat_sanasi = datetime.date.today().strftime('%d.%m.%Y')

    mapping = {
        'BANK_NOMI': settings.get('bank_nomi', ''),
        'BANK_QISQA_NOMI': settings.get('bank_qisqa_nomi', ''),
        'BANK_MANZIL': settings.get('bank_manzil', ''),
        'BANK_EMAIL': settings.get('bank_email', ''),
        'BANK_SAYT': settings.get('bank_sayt', ''),
        'BANK_TEL': settings.get('bank_tel', ''),
        'BANK_MOBIL_ILOVA': settings.get('bank_mobil_ilova', ''),
        'BANK_KODI': settings.get('bank_kodi', ''),
        'ALOQA_MARKAZI_TEL': settings.get('aloqa_markazi_tel', ''),
        'FILIAL_NOMI': settings.get('filial_nomi', ''),
        'FILIAL_TEL': settings.get('filial_tel', ''),
        'RAHBAR_ISM': rahbar_ism or settings.get('rahbar_ism', ''),

        'MIJOZ_ISM': mijoz_ism or '',
        'MIJOZ_MANZIL': mijoz_manzil or '',
        'SARLAVHA': sarlavha,

        'KREDIT_SUMMA': _fmt_summa(portfel_row.get('ead', 0)),
        'KREDIT_MAQSAD': portfel_row.get('tulov_maqsadi', '') or '',
        'HOLAT_SANASI': holat_sanasi,
        'JAMI_QARZ': _fmt_summa(portfel_row.get('jami_qarz', 0)),
        'ASOSIY_QARZ': _fmt_summa(portfel_row.get('asosiy_qarz', 0)),
        'FOIZ_QARZ': _fmt_summa(portfel_row.get('foiz_qarz', 0)),
        'JARIMA': _fmt_summa(portfel_row.get('jarima', 0)),
        'TOLOV_MUDDATI': f"{settings.get('tolov_muddati_kun', '10')} банк иш куни",
        'ANKETA_RAQAM': anketa_raqami or portfel_row.get('anketa_raqami', ''),
    }

    _replace_everywhere(doc, mapping)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def safe_filename(text):
    text = re.sub(r'[\\/*?:"<>|]', '', str(text))
    text = text.strip().replace(' ', '_')
    return text[:80]
